#!/usr/bin/env python
# -*- coding:utf-8 -*-
from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django import conf
from repository import models
import datetime, re, os, time, requests, json, urllib
import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import encryption


def GenerateWordDocument(request, obj_id):
	ret = {"status": "seccuss", "status_code": "200"}

	session = requests.get(url="http://"+ request.get_host() + "/api/resume/candidate/"+ obj_id)
	data = json.loads(session.text)
	newFileName = os.path.splitext(os.path.basename(data["zh_filename"][0]["name"]))[0] + ".docx"
	ResumeTemplatePath = os.path.join(conf.settings.BASE_DIR, "static/firmware/temporary")

	isExists = os.path.exists(ResumeTemplatePath)
	if not isExists: os.makedirs(ResumeTemplatePath)

	obj = models.PreferreResumeTemplate.objects.all().last()
	if obj.resume_template:
		ResumeTemplateUrl = os.path.join(conf.settings.NGINX_MIRROR_ADDRESS, obj.resume_template.url)
		ResumeTemplateName = os.path.basename(ResumeTemplateUrl)
		ResumeTemplateFilePath = os.path.join(ResumeTemplatePath, ResumeTemplateName)
		urllib.request.urlretrieve(ResumeTemplateUrl, ResumeTemplateFilePath)
	else:
		ret["status"] = "failed"
		ret["status_code"] = "404"
		ret["describe"] = "不存在的简历模板,请联系管理员解决此问题!"
		return JsonResponse(ret)

	newWordFilePath = os.path.join(ResumeTemplatePath, newFileName)
	relativePath = os.path.join("static/firmware/temporary", newFileName)

	if not os.path.isfile(newWordFilePath):
		doc = Document(ResumeTemplateFilePath)
		doc.styles['Normal'].font.name = '微软雅黑'
		doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

		# paragraph = doc.add_paragraph("", "Heading 2")
		paragraph = doc.add_paragraph()
		paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		run = paragraph.add_run('个人简历')
		run.bold = True
		run.font.name=u'微软雅黑'
		run.font.size=Pt(11)


		# 换行
		doc.add_paragraph()

		writeData = {}

		# 基础信息
		writeData["BaseInfo"] = {}

		# 自我评价
		writeData["SelfAssessment"] = []

		# 工作经验
		writeData["WorkExperience"] = []

		# 项目经验
		writeData["ProjectExperience"] = []

		# 教育经验
		writeData["EducationExperience"] = []

		# 总和能力
		writeData["ComprehensiveAbility"] = []

		for k,v in data.items():
			if k == "username":
				writeData["BaseInfo"][u"姓名"] = str(v)
			if k == "gender":
				if v == True:
					writeData["BaseInfo"][u"性别"] = u"男"
				else:
					writeData["BaseInfo"][u"性别"] = u"女"
			if k == "birthday":
				writeData["BaseInfo"][u"出生日期"] = str(v)
			if k == "degree":
				writeData["BaseInfo"][u"最高学历"] = str(v)

			if type(v) == list:

				if k == "work_info":
					writeData["WorkExperience"] = v[0]["describe"]

				if k == "project_info":
					writeData["ProjectExperience"] = v[0]["describe"]

				if k == "personal_assessment":
					writeData["SelfAssessment"] = v[0]["describe"]

				if k == "education_info":
					writeData["EducationExperience"] = v[0]["describe"]

				if k == "comprehensive_ability":
					writeData["ComprehensiveAbility"] = v[0]["describe"]

		# 写入基础数据
		paragraph = doc.add_paragraph()
		run = paragraph.add_run(u"基础信息")
		run.font.name='微软雅黑'
		run.font.size=Pt(11)
		run.bold = True

		for k,v in writeData["BaseInfo"].items():
			if v:
				paragraph = doc.add_paragraph()
				run = paragraph.add_run(k + ": " + str(v) if v else "")
				run.font.name='微软雅黑'
				run.font.size=Pt(10.5)
		# 换行
		doc.add_paragraph()
		# print(writeData)

		# 写入自我评价
		if writeData["SelfAssessment"]:

			paragraph = doc.add_paragraph()
			run = paragraph.add_run(u"个人介绍")
			run.bold = True
			run.font.name='微软雅黑'
			run.font.size=Pt(11)

			paragraph = doc.add_paragraph()
			_Data = writeData["SelfAssessment"].replace("\r", '').replace("\n\n", '\n')
			run = paragraph.add_run(_Data)
			run.font.name='微软雅黑'
			run.font.size=Pt(10.5)
		# 换行
		doc.add_paragraph()

		# 写入工作经验
		if writeData["WorkExperience"]:

			paragraph = doc.add_paragraph()
			run = paragraph.add_run(u"工作经验")
			run.bold = True
			run.font.name='微软雅黑'
			run.font.size=Pt(11)

			paragraph = doc.add_paragraph()
			_Data = writeData["WorkExperience"].replace("\r", '').replace("\n\n", '\n')
			run = paragraph.add_run(_Data)
			run.font.name='微软雅黑'
			run.font.size=Pt(10.5)
		# 换行
		doc.add_paragraph()

		# 写入项目经验
		if writeData["ProjectExperience"]:

			paragraph = doc.add_paragraph()
			run = paragraph.add_run(u"项目经验")
			run.bold = True
			run.font.name='微软雅黑'
			run.font.size=Pt(11)

			paragraph = doc.add_paragraph()
			_Data = writeData["ProjectExperience"].replace("\r", '').replace("\n\n", '\n')
			run = paragraph.add_run(_Data)
			run.font.name='微软雅黑'
			run.font.size=Pt(10.5)
		# 换行
		doc.add_paragraph()

		# 写入教育经验
		if writeData["EducationExperience"]:

			paragraph = doc.add_paragraph()
			run = paragraph.add_run(u"教育经验")
			run.bold = True
			run.font.name='微软雅黑'
			run.font.size=Pt(11)

			paragraph = doc.add_paragraph()
			_Data = writeData["EducationExperience"].replace("\r", '').replace("\n\n", '\n')
			run = paragraph.add_run(_Data)
			run.font.name='微软雅黑'
			run.font.size=Pt(10.5)
		# 换行
		doc.add_paragraph()

		# 写入总和能力
		if writeData["ComprehensiveAbility"]:

			paragraph = doc.add_paragraph()
			run = paragraph.add_run(u"总和能力")
			run.bold = True
			run.font.name='微软雅黑'
			run.font.size=Pt(11)

			paragraph = doc.add_paragraph()
			_Data = writeData["ComprehensiveAbility"].replace("\r", '').replace("\n\n", '\n')
			run = paragraph.add_run(_Data)
			run.font.name='微软雅黑'
			run.font.size=Pt(10.5)

		# 保存文件
		doc.save(newWordFilePath)
	currentTime = str(int(time.time()))
	_token = encryption.md5(relativePath + "|" + currentTime) + "|" + currentTime
	ret = {"status": "seccuss", "status_code": "200", "file_path": relativePath, "token": _token}

	# 记录用户下载事件
	models.StatisticalDownloadResume.objects.create(user=request.user, resume=models.ResumeInfo.objects.get(id=int(obj_id)), download_type=2)
	return JsonResponse(ret)